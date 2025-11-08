import os
import re
import uuid
import datetime
import logging
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt as DocxPt
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt as DocxPt
from docx.oxml.shared import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from io import BytesIO
import requests
from utils.file_treatment import _generate_unique_folder, _generate_filename, _public_url, search_image

def create_word(content: list[dict] | str, filename: str, folder_path: str | None = None, title: str | None = None, docx_template_path: str | None = None) -> dict:
    logging.debug("Creating Word document")

    if isinstance(content, str):
        content = _convert_markdown_to_structured(content)
    elif not isinstance(content, list):
        content = []

    if folder_path is None:
        folder_path = _generate_unique_folder()
    if filename:
        filepath = os.path.join(folder_path, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        fname = filename
    else:
        filepath, fname = _generate_filename(folder_path, "docx")

    use_template = False
    doc = None

    if docx_template_path and os.path.exists(docx_template_path):
        try:
            doc = Document(docx_template_path)
            use_template = True
            logging.debug("Using DOCX template")
            # Clear existing body content while retaining styles/theme
            for element in list(doc.element.body):
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    doc.element.body.remove(element)

        except Exception as e:
            logging.warning(f"Failed to load DOCX template: {e}")
            use_template = False
            doc = None

    if not use_template:
        doc = Document()
        logging.debug("Creating new Word document without template")

    if title:
        title_paragraph = doc.add_paragraph(title)
        try:
            title_paragraph.style = doc.styles['Title']
        except KeyError:
            try:
                title_paragraph.style = doc.styles['Heading 1']
            except KeyError:
                run = title_paragraph.runs[0] if title_paragraph.runs else title_paragraph.add_run()
                run.font.size = DocxPt(20)
                run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logging.debug("Document title added")

    for item in content or []:
        if isinstance(item, str):
            doc.add_paragraph(item)
        elif isinstance(item, dict):
            if item.get("type") == "image_query":
                new_item = {
                    "type": "image",
                    "query": item.get("query")
                }
                image_query = new_item.get("query")
                if image_query:
                    logging.debug(f"Image search for the query : {image_query}")
                    image_url = search_image(image_query)
                    if image_url:
                        response = requests.get(image_url)
                        image_data = BytesIO(response.content)
                        doc.add_picture(image_data, width=Inches(6))
                        logging.debug("Image successfully added")
                    else:
                        logging.warning(f"Image search for : '{image_query}'")
            elif "type" in item:
                item_type = item.get("type")
                if item_type == "title":
                    paragraph = doc.add_paragraph(item.get("text", ""))
                    try:
                        paragraph.style = doc.styles['Heading 1']
                    except KeyError:
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.font.size = DocxPt(18)
                        run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logging.debug("Title added")
                elif item_type == "subtitle":
                    paragraph = doc.add_paragraph(item.get("text", ""))
                    try:
                        paragraph.style = doc.styles['Heading 2']
                    except KeyError:
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.font.size = DocxPt(16)
                        run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logging.debug("Subtitle added")
                elif item_type == "paragraph":
                    doc.add_paragraph(item.get("text", ""))
                    logging.debug("Paragraph added")
                elif item_type == "list":
                    items = item.get("items", [])
                    for i, item_text in enumerate(items):
                        paragraph = doc.add_paragraph(item_text)
                        try:
                            paragraph.style = doc.styles['List Bullet']
                        except KeyError:
                            paragraph.style = doc.styles['Normal']
                    logging.debug("List added")
                elif item_type == "image":
                    image_query = item.get("query")
                    if image_query:
                        logging.debug(f"Image search for the query : {image_query}")
                        image_url = search_image(image_query)
                        if image_url:
                            response = requests.get(image_url)
                            image_data = BytesIO(response.content)
                            doc.add_picture(image_data, width=Inches(6))
                            logging.debug("Image successfully added")
                        else:
                            logging.warning(f"Image search for : '{image_query}'")
                elif item_type == "table":
                    data = item.get("data", [])
                    if data:
                        template_table_style = None
                        if use_template:
                            try:
                                for table in doc.tables:
                                    if table.style:
                                        template_table_style = table.style
                                        break
                            except Exception:
                                pass
                        
                        table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
                        
                        if template_table_style:
                            try:
                                table.style = template_table_style
                                logging.debug(f"Applied template table style: {template_table_style.name}")
                            except Exception as e:
                                logging.debug(f"Could not apply template table style: {e}")
                        else:
                            try:
                                for style_name in ['Table Grid', 'Light Grid Accent 1', 'Medium Grid 1 Accent 1', 'Light List Accent 1']:
                                    try:
                                        table.style = doc.styles[style_name]
                                        logging.debug(f"Applied built-in table style: {style_name}")
                                        break
                                    except KeyError:
                                        continue
                            except Exception as e:
                                logging.debug(f"Could not apply any table style: {e}")
                        
                        for i, row in enumerate(data):
                            for j, cell in enumerate(row):
                                cell_obj = table.cell(i, j)
                                cell_obj.text = str(cell)
      	      	      
                                if i == 0:
                                    for paragraph in cell_obj.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
      	      	      
                        if not template_table_style:
                            try:
                                tbl = table._tbl
                                tblPr = tbl.tblPr
                                tblBorders = parse_xml(r'<w:tblBorders {}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tblBorders>'.format(nsdecls('w')))
                                tblPr.append(tblBorders)
                            except Exception as e:
                                logging.debug(f"Could not add table borders: {e}")
                        
                        logging.debug("Table added with improved styling")
            elif "text" in item:
                doc.add_paragraph(item["text"])
                logging.debug("Paragraph added")
    
    doc.save(filepath)
    return {"url": _public_url(folder_path, fname), "path": filepath}

def _convert_markdown_to_structured(markdown_content):
    """
    Converts Markdown content into a structured format for Word
    
    Args:
        markdown_content (str): Markdown content
        
    Returns:
        list: List of objects with 'text' and 'type'
    """
    if not markdown_content or not isinstance(markdown_content, str):
        return []
    
    lines = markdown_content.split('\n')
    structured = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            structured.append({"text": line[2:].strip(), "type": "title"})
        elif line.startswith('## '):
            structured.append({"text": line[3:].strip(), "type": "heading"})
        elif line.startswith('### '):
            structured.append({"text": line[4:].strip(), "type": "subheading"})
        elif line.startswith('#### '):
            structured.append({"text": line[5:].strip(), "type": "subheading"})
        elif line.startswith('- '):
            structured.append({"text": line[2:].strip(), "type": "bullet"})
        elif line.startswith('* '):
            structured.append({"text": line[2:].strip(), "type": "bullet"})
        elif line.startswith('**') and line.endswith('**'):
            structured.append({"text": line[2:-2].strip(), "type": "bold"})
        else:
            structured.append({"text": line, "type": "paragraph"})
    
    return structured

def _apply_text_to_paragraph(para, new_text):
    """
    Apply new text to a paragraph while preserving formatting.
    """
    original_style = para.style
    original_alignment = para.alignment
    
    original_run_format = None
    if para.runs:
        first_run = para.runs[0]
        original_run_format = {
            "font_name": first_run.font.name,
            "font_size": first_run.font.size,
            "bold": first_run.font.bold,
            "italic": first_run.font.italic,
            "underline": first_run.font.underline,
            "color": first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
        }
    
    for _ in range(len(para.runs)):
        para._element.remove(para.runs[0]._element)
    
    if isinstance(new_text, list):
        for i, text_item in enumerate(new_text):
            if i > 0:
                para.add_run("\n")
            run = para.add_run(str(text_item))
            if original_run_format:
                _apply_run_formatting(run, original_run_format)
    else:
        run = para.add_run(str(new_text))
        if original_run_format:
            _apply_run_formatting(run, original_run_format)
    
    if original_style:
        try:
            para.style = original_style
        except Exception:
            pass
    if original_alignment is not None:
        try:
            para.alignment = original_alignment
        except Exception:
            pass


def _apply_run_formatting(run, format_dict):
    """
    Apply formatting from a dict to a run.
    """
    try:
        if format_dict.get("font_name"):
            run.font.name = format_dict["font_name"]
    except Exception:
        pass
    
    try:
        if format_dict.get("font_size"):
            run.font.size = format_dict["font_size"]
    except Exception:
        pass
    
    try:
        if format_dict.get("bold") is not None:
            run.font.bold = format_dict["bold"]
    except Exception:
        pass
    
    try:
        if format_dict.get("italic") is not None:
            run.font.italic = format_dict["italic"]
    except Exception:
        pass
    
    try:
        if format_dict.get("underline") is not None:
            run.font.underline = format_dict["underline"]
    except Exception:
        pass
    
    try:
        if format_dict.get("color"):
            from docx.shared import RGBColor
            run.font.color.rgb = format_dict["color"]
    except Exception:
        pass

def _extract_paragraph_style_info(para):
    """Extract detailed style information from a paragraph"""
    if not para.runs:
        return {}
    
    first_run = para.runs[0]
    return {
        "font_name": first_run.font.name,
        "font_size": first_run.font.size,
        "bold": first_run.font.bold,
        "italic": first_run.font.italic,
        "underline": first_run.font.underline,
        "color": first_run.font.color.rgb if first_run.font.color else None
    }

def _snapshot_runs(p):
    """Return a list of {'text': str, 'font': {...}} for each run in a paragraph."""
    runs = []
    for r in p.runs:
        f = r.font
        font_spec = {
            "name": f.name,
            "size": f.size,
            "bold": f.bold,
            "italic": f.italic,
            "underline": f.underline,
            "color_rgb": getattr(getattr(f.color, "rgb", None), "rgb", None) or getattr(f.color, "rgb", None)
        }
        runs.append({"text": r.text or "", "font": font_spec})
    return runs

def _apply_font(run, font_spec):
    """Apply font specifications to a run."""
    if not font_spec:
        return
    f = run.font
    try:
        if font_spec.get("name") is not None:
            f.name = font_spec["name"]
        if font_spec.get("size") is not None:
            f.size = font_spec["size"]
        if font_spec.get("bold") is not None:
            f.bold = font_spec["bold"]
        if font_spec.get("italic") is not None:
            f.italic = font_spec["italic"]
        if font_spec.get("underline") is not None:
            f.underline = font_spec["underline"]
        rgb = font_spec.get("color_rgb")
        if rgb is not None:
            try:
                f.color.rgb = rgb
            except Exception:
                pass
    except Exception:
        pass
